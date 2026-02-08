"""
Document Parser Module - Multi-format file parsing for RAG ingestion

Supports:
- PDF: via PyMuPDF (fitz) or pdfplumber
- DOCX: via python-docx
- Markdown: Enhanced parsing with section hierarchy
- Plain text: .txt, .py, .js, etc.

Each parser returns extracted text with optional metadata.
"""

from __future__ import annotations
import re
from pathlib import Path
from typing import Dict, Callable, List, Any
from dataclasses import dataclass, field

# ==============================================================================
# DATA STRUCTURES
# ==============================================================================


@dataclass
class ParseResult:
    """Result of parsing a document."""

    content: str
    metadata: Dict[str, Any] = field(default_factory=dict)
    sections: List[Dict[str, Any]] = field(default_factory=list)
    page_count: int = 1
    word_count: int = 0


# ==============================================================================
# PARSER IMPLEMENTATIONS
# ==============================================================================


def parse_plain_text(path: Path) -> ParseResult:
    """Parse plain text files (.txt, .py, .js, etc.)."""
    try:
        content = path.read_text(encoding="utf-8", errors="ignore")
    except Exception as e:
        return ParseResult(content="", metadata={"error": str(e)})

    return ParseResult(
        content=content,
        metadata={"format": "plain_text", "extension": path.suffix},
        word_count=len(content.split()),
    )


def parse_markdown(path: Path) -> ParseResult:
    """
    Parse Markdown files with section hierarchy extraction.

    Extracts:
    - Headers and their hierarchy
    - Code blocks
    - Lists
    - Section structure
    """
    try:
        content = path.read_text(encoding="utf-8", errors="ignore")
    except Exception as e:
        return ParseResult(content="", metadata={"error": str(e)})

    sections = []

    lines = content.split("\n")
    code_block_open = False

    for i, line in enumerate(lines):
        # Track code blocks
        if line.strip().startswith("```"):
            code_block_open = not code_block_open
            continue

        if code_block_open:
            continue

        # Detect headers
        header_match = re.match(r"^(#{1,6})\s+(.+)$", line)
        if header_match:
            level = len(header_match.group(1))
            title = header_match.group(2).strip()
            sections.append(
                {
                    "level": level,
                    "title": title,
                    "line": i + 1,
                }
            )

    # Extract code blocks for metadata
    code_blocks = re.findall(r"```(\w*)\n(.*?)```", content, re.DOTALL)
    languages = list(set(lang for lang, _ in code_blocks if lang))

    return ParseResult(
        content=content,
        metadata={
            "format": "markdown",
            "extension": ".md",
            "header_count": len(sections),
            "code_languages": languages,
            "has_code": len(code_blocks) > 0,
        },
        sections=sections,
        word_count=len(content.split()),
    )


def parse_pdf(path: Path) -> ParseResult:
    """
    Parse PDF files using PyMuPDF (fitz) or pdfplumber.

    Preserves page numbers in metadata.
    """
    content_parts = []
    page_texts = []
    page_count = 0
    metadata = {"format": "pdf", "extension": ".pdf"}

    # Try PyMuPDF first (faster)
    try:
        import fitz  # PyMuPDF

        doc = fitz.open(str(path))
        page_count = len(doc)

        for page_num, page in enumerate(doc, 1):
            text = page.get_text()
            if text.strip():
                page_texts.append(
                    {
                        "page": page_num,
                        "text": text,
                        "char_count": len(text),
                    }
                )
                content_parts.append(f"[Page {page_num}]\n{text}")

        doc.close()
        metadata["parser"] = "pymupdf"

    except ImportError:
        # Fallback to pdfplumber
        try:
            import pdfplumber

            with pdfplumber.open(str(path)) as pdf:
                page_count = len(pdf.pages)

                for page_num, page in enumerate(pdf.pages, 1):
                    text = page.extract_text() or ""
                    if text.strip():
                        page_texts.append(
                            {
                                "page": page_num,
                                "text": text,
                                "char_count": len(text),
                            }
                        )
                        content_parts.append(f"[Page {page_num}]\n{text}")

            metadata["parser"] = "pdfplumber"

        except ImportError:
            return ParseResult(
                content="",
                metadata={
                    "format": "pdf",
                    "error": "No PDF parser available. Install pymupdf or pdfplumber.",
                },
            )
        except Exception as e:
            return ParseResult(
                content="",
                metadata={"format": "pdf", "error": str(e)},
            )
    except Exception as e:
        return ParseResult(
            content="",
            metadata={"format": "pdf", "error": str(e)},
        )

    content = "\n\n".join(content_parts)
    metadata["page_count"] = page_count
    metadata["pages_with_text"] = len(page_texts)

    return ParseResult(
        content=content,
        metadata=metadata,
        page_count=page_count,
        word_count=len(content.split()),
    )


def parse_docx(path: Path) -> ParseResult:
    """
    Parse DOCX files using python-docx.

    Extracts text from paragraphs and tables.
    """
    content_parts = []
    metadata = {"format": "docx", "extension": ".docx"}

    try:
        from docx import Document

        doc = Document(str(path))

        # Extract paragraphs
        paragraph_count = 0
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                content_parts.append(text)
                paragraph_count += 1

        # Extract tables
        table_count = 0
        for table in doc.tables:
            table_rows = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                table_rows.append(" | ".join(cells))

            if table_rows:
                content_parts.append("\n[Table]\n" + "\n".join(table_rows))
                table_count += 1

        metadata["paragraph_count"] = paragraph_count
        metadata["table_count"] = table_count
        metadata["parser"] = "python-docx"

    except ImportError:
        return ParseResult(
            content="",
            metadata={
                "format": "docx",
                "error": "python-docx not installed. Run: pip install python-docx",
            },
        )
    except Exception as e:
        return ParseResult(
            content="",
            metadata={"format": "docx", "error": str(e)},
        )

    content = "\n\n".join(content_parts)

    return ParseResult(
        content=content,
        metadata=metadata,
        word_count=len(content.split()),
    )


def parse_html(path: Path) -> ParseResult:
    """
    Parse HTML files - extract text content.

    Simple regex-based extraction (no external dependencies).
    """
    try:
        raw = path.read_text(encoding="utf-8", errors="ignore")
    except Exception as e:
        return ParseResult(content="", metadata={"error": str(e)})

    # Remove script and style blocks
    content = re.sub(r"<script[^>]*>.*?</script>", "", raw, flags=re.DOTALL | re.IGNORECASE)
    content = re.sub(r"<style[^>]*>.*?</style>", "", content, flags=re.DOTALL | re.IGNORECASE)

    # Remove HTML tags
    content = re.sub(r"<[^>]+>", " ", content)

    # Decode common entities
    content = content.replace("&nbsp;", " ")
    content = content.replace("&amp;", "&")
    content = content.replace("&lt;", "<")
    content = content.replace("&gt;", ">")
    content = content.replace("&quot;", '"')

    # Normalize whitespace
    content = " ".join(content.split())

    return ParseResult(
        content=content,
        metadata={"format": "html", "extension": path.suffix},
        word_count=len(content.split()),
    )


def parse_json(path: Path) -> ParseResult:
    """Parse JSON files - format nicely for ingestion."""
    try:
        import json

        raw = path.read_text(encoding="utf-8", errors="ignore")
        data = json.loads(raw)
        content = json.dumps(data, indent=2, ensure_ascii=False)
    except Exception as e:
        return ParseResult(content="", metadata={"error": str(e)})

    return ParseResult(
        content=content,
        metadata={"format": "json", "extension": ".json"},
        word_count=len(content.split()),
    )


def parse_yaml(path: Path) -> ParseResult:
    """Parse YAML files."""
    try:
        content = path.read_text(encoding="utf-8", errors="ignore")
    except Exception as e:
        return ParseResult(content="", metadata={"error": str(e)})

    return ParseResult(
        content=content,
        metadata={"format": "yaml", "extension": path.suffix},
        word_count=len(content.split()),
    )


# ==============================================================================
# PARSER REGISTRY
# ==============================================================================

PARSERS: Dict[str, Callable[[Path], ParseResult]] = {
    # Plain text
    ".txt": parse_plain_text,
    ".log": parse_plain_text,
    ".csv": parse_plain_text,
    # Code files
    ".py": parse_plain_text,
    ".js": parse_plain_text,
    ".ts": parse_plain_text,
    ".jsx": parse_plain_text,
    ".tsx": parse_plain_text,
    ".java": parse_plain_text,
    ".c": parse_plain_text,
    ".cpp": parse_plain_text,
    ".h": parse_plain_text,
    ".hpp": parse_plain_text,
    ".rs": parse_plain_text,
    ".go": parse_plain_text,
    ".rb": parse_plain_text,
    ".php": parse_plain_text,
    ".cs": parse_plain_text,
    ".swift": parse_plain_text,
    ".kt": parse_plain_text,
    ".scala": parse_plain_text,
    ".sql": parse_plain_text,
    ".sh": parse_plain_text,
    ".bash": parse_plain_text,
    ".ps1": parse_plain_text,
    ".bat": parse_plain_text,
    # Markdown
    ".md": parse_markdown,
    ".markdown": parse_markdown,
    ".mdx": parse_markdown,
    # Documents
    ".pdf": parse_pdf,
    ".docx": parse_docx,
    # Web
    ".html": parse_html,
    ".htm": parse_html,
    ".xml": parse_plain_text,
    # Data formats
    ".json": parse_json,
    ".yaml": parse_yaml,
    ".yml": parse_yaml,
    ".toml": parse_plain_text,
    ".ini": parse_plain_text,
    ".cfg": parse_plain_text,
    ".conf": parse_plain_text,
}


# ==============================================================================
# PUBLIC API
# ==============================================================================


def parse_file(path: Path, fallback_to_text: bool = True) -> ParseResult:
    """
    Parse a file using the appropriate parser based on extension.

    Args:
        path: Path to the file to parse.
        fallback_to_text: If True, try plain text parsing for unknown extensions.

    Returns:
        ParseResult with content and metadata.
    """
    if not path.exists():
        return ParseResult(
            content="",
            metadata={"error": f"File not found: {path}"},
        )

    ext = path.suffix.lower()

    # Get parser from registry
    parser = PARSERS.get(ext)

    if parser:
        return parser(path)
    elif fallback_to_text:
        # Try plain text for unknown extensions
        try:
            return parse_plain_text(path)
        except Exception:
            return ParseResult(
                content="",
                metadata={"error": f"Unable to parse file with extension: {ext}"},
            )
    else:
        return ParseResult(
            content="",
            metadata={"error": f"No parser registered for extension: {ext}"},
        )


def get_supported_extensions() -> List[str]:
    """Return list of supported file extensions."""
    return list(PARSERS.keys())


def register_parser(extension: str, parser: Callable[[Path], ParseResult]) -> None:
    """
    Register a custom parser for a file extension.

    Args:
        extension: File extension (with dot, e.g., ".xyz")
        parser: Parser function that takes Path and returns ParseResult
    """
    PARSERS[extension.lower()] = parser


def can_parse(path: Path) -> bool:
    """Check if a file can be parsed."""
    return path.suffix.lower() in PARSERS
